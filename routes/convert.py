from flask import Blueprint, request, render_template
from pdf2image import convert_from_path
from pdf2docx import Converter
from PIL import Image
from pypdf import PdfReader
from werkzeug.utils import secure_filename
import fitz  # (PyMuPDF)
from docx import Document

def normalize_vertical_table(rows):
    """
    Convierte columnas verticales en filas horizontales
    """
    new_rows = []
    temp = []

    for r in rows:
        if len(r) == 1:
            temp.append(r[0])

            # cada 3 elementos → crear fila
            if len(temp) == 3:
                new_rows.append(temp)
                temp = []
        else:
            if temp:
                new_rows.append(temp)
                temp = []
            new_rows.append(r)

    if temp:
        new_rows.append(temp)

    return new_rows


def extract_tables_to_docx(input_pdf, output_docx):
    import fitz
    from docx import Document

    doc = fitz.open(input_pdf)
    word = Document()

    for page_num, page in enumerate(doc):
        blocks = page.get_text("blocks")

        blocks = sorted(blocks, key=lambda b: b[1])

        table_data = []

        for b in blocks:
            text = b[4].strip()

            if len(text.split()) > 3:
                row = [cell.strip() for cell in text.split()]
                table_data.append(row)

            else:
                if table_data:
                    # 🔥 AQUÍ SE APLICA LA MAGIA
                    table_data = normalize_vertical_table(table_data)

                    rows = len(table_data)
                    cols = max(len(r) for r in table_data)

                    table = word.add_table(rows=rows, cols=cols)
                    style_table_full_width(table)

                    for i, row in enumerate(table_data):
                        for j, val in enumerate(row):
                            table.cell(i, j).text = val

                    table_data = []

                word.add_paragraph(text)

        if table_data:
            # 🔥 TAMBIÉN AQUÍ (IMPORTANTE)
            table_data = normalize_vertical_table(table_data)

            rows = len(table_data)
            cols = max(len(r) for r in table_data)

            table = word.add_table(rows=rows, cols=cols)

            for i, row in enumerate(table_data):
                for j, val in enumerate(row):
                    table.cell(i, j).text = val

    word.save(output_docx)
    doc.close()
    
def style_table_full_width(table):
    from docx.shared import Inches
    from docx.enum.table import WD_TABLE_ALIGNMENT

    # 🔹 centrar tabla
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 🔹 desactivar autoajuste raro
    table.autofit = False

    # 🔹 ancho total aproximado (A4 / carta)
    total_width = Inches(6.5)

    cols = len(table.columns)
    col_width = total_width / cols

    for row in table.rows:
        for cell in row.cells:
            cell.width = col_width

            for paragraph in cell.paragraphs:
                paragraph.alignment = 1

import os, zipfile, uuid, subprocess

def allowed_file(filename, allowed_extensions):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in allowed_extensions
ALLOWED_PDF = {"pdf"}
ALLOWED_IMAGES = {"png", "jpg", "jpeg"}
ALLOWED_WORD = {"docx"}

convert_bp = Blueprint("convert", __name__)

BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")
OUTPUT_FOLDER = os.path.join(BASE_DIR, "outputs")

# PDF TO IMG
@convert_bp.route("/pdf_to_img", methods=["POST"])
def pdf_to_img():
    file = request.files["pdf"]

    if not allowed_file(file.filename, ALLOWED_PDF):
        return render_template("result.html", error="Solo se permiten archivos PDF")
    
    filename_secure = secure_filename(file.filename)
    pdf_path = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4()}_{filename_secure}")
    file.save(pdf_path)

    images = convert_from_path(pdf_path)

    zip_name = f"{uuid.uuid4()}.zip"
    zip_path = os.path.join(OUTPUT_FOLDER, zip_name)

    temp_images = []

    with zipfile.ZipFile(zip_path, "w") as zipf:
        for i, img in enumerate(images):
            img_name = f"{uuid.uuid4()}.png"
            img_path = os.path.join(OUTPUT_FOLDER, img_name)
            img.save(img_path, "PNG")
            zipf.write(img_path, f"page_{i+1}.png")
            temp_images.append(img_path)

    # 🔥 limpiar temporales
    os.remove(pdf_path)
    for img in temp_images:
        os.remove(img)

    return render_template("result.html", filename=zip_name)

# IMG TO PDF
@convert_bp.route("/img_to_pdf", methods=["POST"])
def img_to_pdf():
    files = request.files.getlist("images")
    
    paths = []
    image_list = []

    for file in files:
        if not allowed_file(file.filename, ALLOWED_IMAGES):
            return render_template("result.html", error="Solo imágenes JPG o PNG")
        filename_secure = secure_filename(file.filename)
        path = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4()}_{filename_secure}")
        file.save(path)
        paths.append(path)

        img = Image.open(path).convert("RGB")
        image_list.append(img)

    filename = f"{uuid.uuid4()}.pdf"
    output_path = os.path.join(OUTPUT_FOLDER, filename)

    image_list[0].save(output_path, save_all=True, append_images=image_list[1:])

    for path in paths:
        os.remove(path)

    return render_template("result.html", filename=filename)

# PDF TO WORD (MEJORADO)
@convert_bp.route("/pdf_to_word", methods=["POST"])
def pdf_to_word():
    file = request.files.get("pdf")

    if not allowed_file(file.filename, ALLOWED_PDF):
        return render_template("result.html", error="Solo se permiten archivos PDF")

    filename_secure = secure_filename(file.filename)
    input_path = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4()}_{filename_secure}")
    file.save(input_path)

    reader = PdfReader(input_path)

    if len(reader.pages) > 50:
        os.remove(input_path)
        return render_template("result.html", filename=None, error="Máximo 50 páginas")

    filename = f"{uuid.uuid4()}.docx"
    output_path = os.path.join(OUTPUT_FOLDER, filename)

    # =========================
    # 🔍 ANALIZAR PDF (NUEVO)
    # =========================
    doc = fitz.open(input_path)

    has_text = False
    has_images = False

    for page in doc:
        if page.get_text().strip():
            has_text = True
        if page.get_images():
            has_images = True

    doc.close()

    # =========================
    # 🔀 DECISIÓN INTELIGENTE
    # =========================

    # PDF ESCANEADO (SIN TEXTO)
    if not has_text:
        os.remove(input_path)
        return render_template(
            "result.html",
            error="Este PDF parece escaneado. Usa la herramienta OCR PDF."
        )

    # PDF COMPLEJO (IMÁGENES / DISEÑO)
    if has_images:
        try:
            extract_tables_to_docx(input_path, output_path)
        except Exception as e:
            print("Fallback a pdf2docx:", e)

            cv = Converter(input_path)
            cv.convert(output_path, start=0, end=None)
            cv.close()

    # PDF SIMPLE (como ya lo tenías)
    else:
        cv = Converter(input_path)
        cv.convert(output_path)
        cv.close()

    os.remove(input_path)

    return render_template("result.html", filename=filename)

# WORD TO PDF
@convert_bp.route("/word_to_pdf", methods=["POST"])
def word_to_pdf():
    file = request.files.get("file")

    if not allowed_file(file.filename, ALLOWED_WORD):
        return render_template("result.html", error="Solo archivos Word (.docx)")
    
    filename_secure = secure_filename(file.filename)
    input_path = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4()}_{filename_secure}")
    file.save(input_path)

    subprocess.run(
        [
            "soffice",  # 🔥 esto funciona en Render/Linux
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            OUTPUT_FOLDER,
            input_path,
        ],
        check=True,
    )

    output_name = os.path.splitext(os.path.basename(input_path))[0] + ".pdf"

    os.remove(input_path)

    return render_template("result.html", filename=output_name)
